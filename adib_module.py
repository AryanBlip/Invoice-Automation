from datetime import datetime
from tkinter import filedialog, messagebox, ttk, Label, Frame, Entry
from tkinter import TclError, END, Button, VERTICAL, CENTER 
from docx2pdf import convert
from docx import Document, enum
from docx.shared import Pt
from pandas import read_excel, notna
from re import sub
from num2words import num2words
import sys
from os import path

def resource_path(relative_path):
    if getattr(sys, 'frozen', False):
        base_path = path.dirname(path.abspath(sys.argv[0]))
    else:
        base_path = path.dirname(path.abspath(__file__))
        
    return path.join(base_path, relative_path)

class InvoiceAutomation:
    def __init__(self, parent_window, excel_file_path, main_app_root):
        self.template = resource_path('ADIBtemplate.docx')
        self.root = parent_window
        self.main_app_root = main_app_root
        self.root.title("Invoice Automation - ADIB")
        self.root.geometry("800x600") # Increased size for better table visibility
        
        self.root.protocol("WM_DELETE_WINDOW", self.on_close)

        self.excel_file_path = excel_file_path
        
        # UI Elements
        main_frame = Frame(self.root)
        main_frame.pack(fill="both", expand=True, padx=10, pady=10)

        # --- Input Fields ---
        input_frame = Frame(main_frame)
        input_frame.pack(fill="x", pady=5)
        
        self.invoice_date = datetime.today().strftime('%d/%m/%Y')
        year = self.invoice_date.split("/")[-1]

        self.invoice_number_label = Label(input_frame, text="Invoice Number:", anchor="e")
        self.invoice_number_label.grid(row=0, column=0, padx=5, pady=2, sticky="ew")
        self.invoice_number_entry = Entry(input_frame)
        self.invoice_number_entry.insert(END, self.getNextInvoiceNumber() + f"/{year}")
        self.invoice_number_entry.grid(row=0, column=1, padx=5, pady=2, sticky="ew")


        self.invoice_date_label = Label(input_frame, text="Invoice Date:", anchor="e")
        self.invoice_date_label.grid(row=1, column=0, padx=5, pady=2, sticky="ew")
        self.invoice_date_entry = Entry(input_frame)
        self.invoice_date_entry.insert(END, self.invoice_date)
        self.invoice_date_entry.grid(row=1, column=1, padx=5, pady=2, sticky="ew")


        self.month_year_label = Label(input_frame, text="Month Year:", anchor="e")
        self.month_year_label.grid(row=2, column=0, padx=5, pady=2, sticky="ew")
        self.month_year_entry = Entry(input_frame)
        self.month_year_entry.grid(row=2, column=1, padx=5, pady=2, sticky="ew")
        
        input_frame.columnconfigure(1, weight=1)

        # --- Data Display Table (Treeview) ---
        table_frame = Frame(main_frame)
        table_frame.pack(fill="both", expand=True, pady=10)
        
        # Define columns for the Treeview table
        columns = ("Disbursal Date", "Type", "Customer Name", "Loan Amount", "Payment Slab", "Incentive")
        self.tree = ttk.Treeview(table_frame, columns=columns, show="headings")
        
        for col in columns:
            self.tree.heading(col, text=col)
            self.tree.column(col, width=120, anchor=CENTER)
            
        self.tree.pack(side="left", fill="both", expand=True)

        # Add a vertical scrollbar
        scrollbar = ttk.Scrollbar(table_frame, orient=VERTICAL, command=self.tree.yview)
        self.tree.configure(yscrollcommand=scrollbar.set)
        scrollbar.pack(side="right", fill="y")
        
        # Make the table cells editable on double-click
        self.tree.bind('<Double-1>', self.on_double_click)

        # --- Note below the table ---
        note_label = Label(main_frame, 
                              text="Note: The 'Disbursal Date' will be set to the 'Month Year' entered above when creating the invoice.", 
                              font=('Arial', 9, 'italic'),
                              fg='gray')
        note_label.pack(pady=(5, 0))

        # --- Buttons ---
        button_frame = Frame(main_frame)
        button_frame.pack(fill="x", pady=5)
        
        self.create_button = Button(button_frame, text="Create Invoice", command=self.create_invoice)
        self.create_button.pack(side="left", padx=5, expand=True)

        # Load data into the Treeview
        self.load_data_from_excel()

    @staticmethod
    def getNextInvoiceNumber():
        counterPath = resource_path("counter.txt")
        
        # Read current
        with open(counterPath, "r") as f:
            line = f.readline()
            current_num = line.split(" ")[-1]

        return current_num.strip().zfill(3)
    
    @staticmethod
    def incrementInvoiceCounter():
        counterPath = resource_path("counter.txt")
        # Read current
        with open(counterPath, "r") as f:
            line = f.readline()
            current_num = line.split(" ")[-1]

        # Write next
        next_num = int(current_num) + 1
        with open(counterPath, "w") as f:
            f.write(f"Next Invoice Number : {next_num}")

    @staticmethod
    def IntComma(num):
        parts = num.split('.')
        integer_part = parts[0]
        fractional_part = parts[1] if len(parts) > 1 else None

        # Format the integer part with commas
        n = len(integer_part)
        if n <= 3:
            formatted_integer_part = integer_part
        else:
            result = []
            for i in range(n - 1, -1, -3):
                start = max(0, i - 2)
                result.append(integer_part[start:i + 1])
            formatted_integer_part = ",".join(reversed(result))

        # Reassemble the number
        if fractional_part is not None:
            return f"{formatted_integer_part}.{fractional_part}"
        else:
            return formatted_integer_part

    @staticmethod
    def clean_and_convert_Integer(raw_string):
        # Regex: Find and remove everything that is NOT a digit (\d), a dot (\.), or a minus sign (\-)
        # This single line handles spaces, commas, and \xa0 characters effectively.
        cleaned_string = sub(r'[^\d\.\-]', '', str(raw_string))
        
        # The sub function alone is often enough, but a final .strip() is safe practice
        return str(cleaned_string.strip())
    
    @staticmethod
    def clean_and_convert_String(raw_string):
        # Regex: Find and remove everything that is NOT a digit (\d), a dot (\.), or a minus sign (\-)
        # This single line handles spaces, commas, and \xa0 characters effectively.
        cleaned_string = raw_string.replace('\xa0', '').strip()
        
        # The sub function alone is often enough, but a final .strip() is safe practice
        return str(cleaned_string.strip())

    def load_data_from_excel(self):
        """Loads customer data from the selected Excel file into the Treeview."""
        try:
            # Clear existing data in the treeview
            for item in self.tree.get_children():
                self.tree.delete(item)

            self.true_header = []

            # true header is the fetched header from excel file to identify data
            self.customer_data = read_excel(self.excel_file_path, header=None)
            for index, row_data in self.customer_data.iterrows():
                current_list = [self.clean_and_convert_String(str(x)).lower() for x in row_data]
                if "customer name" in current_list:
                    self.true_header = current_list
                    break
            
            # Iterate through DataFrame and insert into the Treeview
            for index, row_data in self.customer_data.iterrows():
                try:
                    disbursal_date = "(Month Year as stated above)"

                    customer_name_index = self.true_header.index('customer name')
                    customer_name = str(self.clean_and_convert_String(str(row_data.iloc[ customer_name_index ]))).title() if notna(row_data.iloc[ customer_name_index ]) else ""
                    
                    loan_amount_index = self.true_header.index('contract amt')
                    loan_amount = float(self.clean_and_convert_Integer(str(row_data[ loan_amount_index ])))

                    payment_slab = 0.9

                    incentive = loan_amount * payment_slab / 100

                    # Insert the row into the Treeview
                    self.tree.insert("", END, values=(
                        disbursal_date, 
                        "New",
                        customer_name,
                        self.IntComma(f"{loan_amount:.2f}"), 
                        str(payment_slab) + "%", 
                        self.IntComma(f"{incentive:.2f}")
                    ))
                except Exception as e:
                    # TODO: DELETE STMT DURING DEPLOYMENT
                    print(f"Error processing row: {e}")
                    pass
            if not self.tree.get_children():
                raise ValueError("Header must be included in Excel file")

        except FileNotFoundError:
            messagebox.showerror("Error", f"Excel file not found at: {self.excel_file_path}")

        except IndexError:
            messagebox.showerror("Error", "Excel file does not contain the required columns.\nVerify correct File / Format(.xlsx) and it's columns.")
            self.root.destroy()
            self.main_app_root.destroy()

        except ValueError as e:
            messagebox.showerror("Error", f"Verify Correct File / Format(.xlsx) and it's columns.\nError: {e}")
            self.root.destroy()
            self.main_app_root.destroy()

        except TclError as e:
            messagebox.showerror("Error", f"Failed to load Excel data.\nVerify Correct File / Format(.xlsx) and it's columns.\nError: {e}")
            self.root.destroy()
            self.main_app_root.destroy()

        except Exception as e:
            messagebox.showerror("Error", f"Unknown Error: {e}.\nContact the developer (Aryan).")
            
    def on_double_click(self, event):
        # Handles double-click events to make a cell editable
        region = self.tree.identify("region", event.x, event.y)
        if region != "cell":
            return
            
        column = self.tree.identify_column(event.x)
        column_index = int(column[1:]) - 1 # Get 0-indexed column number

        # Do not allow editing for the 'Incentive' column (index 5)
        if column_index == 5:
            messagebox.showinfo("Info", "The Incentive field is automatically calculated and is not editable.")
            return

        item_id = self.tree.identify_row(event.y)
        if not item_id:
            return

        cell_value = self.tree.item(item_id, "values")[column_index]

        # Create a temporary entry widget for editing
        x, y, width, height = self.tree.bbox(item_id, column)
        entry_edit = ttk.Entry(self.tree, justify='center')
        entry_edit.place(x=x, y=y, width=width, height=height)
        entry_edit.insert(0, cell_value)
        entry_edit.focus()

        def save_edit(event):
            new_value = entry_edit.get()
            current_values = list(self.tree.item(item_id, "values"))
            
            # Update the Treeview with the new value
            current_values[column_index] = new_value
            
            # Recalculate incentive if Loan Amount or Payment Slab changed
            if column_index in (3,4):
                try:
                    # Clean the loan amount string before converting to float
                    loan_amount_str = sub(r'[^\d.]', '', str(current_values[3]))
                    loan_amount = float(loan_amount_str)
                    
                    # Clean the payment slab string before converting to float
                    payment_slab_str = str(current_values[4]).strip('%')
                    payment_slab = float(payment_slab_str)
                    
                    incentive = loan_amount * payment_slab / 100
                    current_values[4] += "%" if "%" not in current_values[4] else ""
                    current_values[5] = self.IntComma(f"{incentive:.2f}")

                except ValueError:
                    messagebox.showerror("Invalid Input", "Loan Amount and Payment Slab must be numbers.")
                    return

            self.tree.item(item_id, values=current_values)
            entry_edit.destroy()

        entry_edit.bind("<Return>", save_edit)
        entry_edit.bind("<FocusOut>", save_edit)

    def run(self):
        self.root.mainloop()

    def on_close(self):
        self.root.destroy()
        self.main_app_root.destroy()
    
    @staticmethod
    def replace_text(paragraph, old_text, new_text, isTable=False):
        full = "".join(run.text for run in paragraph.runs) or paragraph.text
        if old_text in full:
            replaced = full.replace(old_text, new_text)
            if paragraph.runs:
                paragraph.runs[0].text = replaced
                for r in paragraph.runs[1:]:
                    r.text = ""
            else:
                paragraph.text = replaced

    def format_table_cells(self, table):
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
                
    @staticmethod
    def number_to_words(num):
        try:
            # Handle integers
            if isinstance(num, int):
                return num2words(num)

            # Handle floats
            elif isinstance(num, float):
                integer_part = int(num)
                decimal_part = str(num).split(".")[1]

                # Convert integer part
                wordsFirst = ''.join(num2words(integer_part).split(" and"))
                wordsSecond = num2words(decimal_part)

                return f"{wordsFirst}".title() + " and " + f"{wordsSecond} Fills Only".title()

            else:
                return "Invalid Number"

        except Exception as e:
            return f"Error: {e}"
        
    @staticmethod
    def convertToFull(month_year):
        try:
            
            if " " in month_year:
                month_year = month_year.split(" ")
                month = month_year[0]
                year = month_year[1]

            else :
                month = ""
                other = ""
                year = ""

                for i in month_year:
                    if i.isalpha():
                        month += i
                    elif i.isnumeric():
                        year += i
                    else:
                        other += i

            months = {
                "jan" : "january", 
                "feb" : "february",
                "mar" : "march",
                "apr" : "april",
                "may" : "may",
                "jun" : "june",
                "jul" : "july",
                "aug" : "august",
                "sep" : "september",
                "oct" : "october",
                "nov" : "november",
                "dec" : "december"
            }

            full_month_year = months.get(month, month) + ((" " + year) if year else year)

            months = {
                "january" : "jan", 
                "february" : "feb",
                "march" : "mar",
                "april" : "apr",
                "may" : "may",
                "june" : "jun",
                "july" : "jul",
                "august" : "aug",
                "september" : "sep",
                "october" : "oct",
                "november" : "nov",
                "december" : "dec"
            }

            half_month_year = months.get(month, month) + ((" " + year) if year else year)

            return (half_month_year, full_month_year)

        except IndexError:
            messagebox.showerror("Missing Information", "Please enter month and year")

    def create_invoice(self):
        # Check if invoice number is empty
        if not self.invoice_number_entry.get():
            messagebox.showerror("Missing Information", "Please enter the Invoice Number.")
            return

        # Check if month/year is empty
        if not self.month_year_entry.get():
            messagebox.showerror("Missing Information", "Please enter the Month Year.")
            return

        # Get data directly from the Treeview
        items = self.tree.get_children()
        if not items:
            messagebox.showerror("Error", "No Excel data loaded. Please select a valid file.")
            return
        
        self.root.withdraw()  # Hide the current window

        totalLoanAmount = 0.0
        totalIncentive = 0.0
        
        doc = Document(self.template)

        replacements = {
            "[date today]" : self.invoice_date_entry.get(),
            "[invoice no]": self.invoice_number_entry.get(),
            "[FullMonth year]" : self.convertToFull( self.month_year_entry.get().lower() )[1].title(),
            "[month year]" : self.convertToFull( self.month_year_entry.get().lower() )[0].title()
        }

        customer_table = None
        for table in doc.tables:
            headers = (cell.text.strip().lower() for cell in table.rows[0].cells)
            if "customer name" in headers:
                customer_table = table
                break

        if not customer_table:
            messagebox.showerror("Error", "Customer table not found in template.")
            return

        # Loop through the rows of the Treeview widget
        for item_id in items:
            row_data = self.tree.item(item_id, 'values')
            try:
                row = customer_table.add_row().cells
                
                # Disbursal Date - Left-aligned by default, no change needed
                row[0].text = self.month_year_entry.get().title() or str(row_data[0])
                
                # Type (New) - Middle-aligned
                row[1].text = str(row_data[1])
                row[1].paragraphs[0].alignment = enum.text.WD_ALIGN_PARAGRAPH.CENTER
                
                # Customer Name - Left-aligned by default, no change needed
                row[2].text = str(row_data[2])
                
                # Loan Amount - Right-aligned
                originalLoanAmount = float(self.clean_and_convert_Integer(str(row_data[3])))
                row[3].text = self.IntComma(f"{originalLoanAmount:.2f}")
                row[3].paragraphs[0].alignment = enum.text.WD_ALIGN_PARAGRAPH.RIGHT
                totalLoanAmount += originalLoanAmount

                # Payment Slab - Middle-aligned
                row[4].text = str(row_data[4])
                row[4].paragraphs[0].alignment = enum.text.WD_ALIGN_PARAGRAPH.CENTER
                
                # Incentive - Right-aligned
                currentIncentive = float(self.clean_and_convert_Integer(str(row_data[5])))
                row[5].text = self.IntComma(f"{currentIncentive:.2f}")
                row[5].paragraphs[0].alignment = enum.text.WD_ALIGN_PARAGRAPH.RIGHT
                totalIncentive += currentIncentive
                
            except Exception as e:
                messagebox.showerror("Data Error", f"Error processing data for row {row_data[2]}. Check data types. Error: {e}")
                return 
        
        vatIncentive = totalIncentive * 1.05
        vatIncentive = float(f"{vatIncentive:.2f}")
        replacements.update({
            "[five percent]": self.IntComma(f"{totalIncentive * 0.05:.2f}"),
            "[total loan]": self.IntComma(f"{totalLoanAmount:.2f}"),
            "[total incent]": self.IntComma(f"{totalIncentive:.2f}"),
            "[VAT&incent]": self.IntComma(f"{vatIncentive}"),
            "[AmtinWords]" : f"{self.number_to_words(vatIncentive)}"
        })

        self.format_table_cells(customer_table)

        for paragraph in doc.paragraphs:
            for old_text, new_text in replacements.items():
                self.replace_text(paragraph, old_text, new_text)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for old_text, new_text in replacements.items():
                            self.replace_text(paragraph, old_text, new_text)

        try:
            save_path = filedialog.asksaveasfilename(defaultextension=".pdf", filetypes=[("PDF Documents", "*.pdf")])
            if save_path:
                doc.save('filled.docx')
                convert('filled.docx', save_path)
                print("\nINVOICE HAS BEEN GENERATED !")

                # INCREMENT INVOICE NUMBER COUNTER.TXT
                self.incrementInvoiceCounter()
                messagebox.showinfo("Success", "Invoice created and saved successfully!")
            
            try:
                self.root.destroy()
                self.main_app_root.destroy()
            except TclError:
                # THE WINDOW HAS ALREADY BEEN DESTROYED
                pass

            # IGNORE ANY OTHER EXCEPTION
            except Exception as e:
                pass
        except PermissionError:
            messagebox.showerror("Error", "Permission denied. Please close the file and try again.")