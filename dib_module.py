from datetime import datetime
from tkinter import filedialog, messagebox, ttk, Label, Frame, Entry
from tkinter import TclError, END, Button, VERTICAL, CENTER 
from docx2pdf import convert
from docx import Document, enum
from docx.shared import Pt
from pandas import read_excel, notna
from re import sub
from num2words import num2words
import sys
from os import path

def resource_path(relative_path):
    """Get path to resource for both .py (VS Code) and Nuitka .exe"""
    if getattr(sys, 'frozen', False):
        # Nuitka (compiled exe)
        base_path = path.dirname(sys.executable)
    else:
        # Normal Python (.py script in VS Code)
        base_path = path.dirname(__file__)
    return path.join(base_path, relative_path)

class InvoiceAutomation:
    def __init__(self, parent_window, excel_file_path, main_app_root):
        self.template = resource_path('DIBtemplate.docx')
        self.root = parent_window
        self.main_app_root = main_app_root
        self.root.title("Invoice Automation - DIB")
        self.root.geometry("900x600")  # Bigger window for DIB since more columns

        self.root.protocol("WM_DELETE_WINDOW", self.on_close)

        self.excel_file_path = excel_file_path

        # --- UI Layout ---
        main_frame = Frame(self.root)
        main_frame.pack(fill="both", expand=True, padx=10, pady=10)

        # --- Input Fields ---
        input_frame = Frame(main_frame)
        input_frame.pack(fill="x", pady=5)

        self.invoice_number_label = Label(input_frame, text="Invoice Number:", anchor="e")
        self.invoice_number_label.grid(row=0, column=0, padx=5, pady=2, sticky="ew")
        self.invoice_number_entry = Entry(input_frame)
        self.invoice_number_entry.grid(row=0, column=1, padx=5, pady=2, sticky="ew")

        self.month_year_label = Label(input_frame, text="Month Year:", anchor="e")
        self.month_year_label.grid(row=1, column=0, padx=5, pady=2, sticky="ew")
        self.month_year_entry = Entry(input_frame)
        self.month_year_entry.grid(row=1, column=1, padx=5, pady=2, sticky="ew")

        input_frame.columnconfigure(1, weight=1)

        # --- Treeview Table ---
        table_frame = Frame(main_frame)
        table_frame.pack(fill="both", expand=True, pady=10)

        columns = ("Disbursal Date", "App reference", "Customer Name",
                   "Loan Amount", "Payment Slab", "Payout",
                   "5% VAT", "Incentive")

        self.tree = ttk.Treeview(table_frame, columns=columns, show="headings")

        for col in columns:
            self.tree.heading(col, text=col)
            self.tree.column(col, width=120, anchor=CENTER)
            
        # Explicitly center the Loan Amount column
        self.tree.column("Loan Amount", anchor=CENTER)

        self.tree.pack(side="left", fill="both", expand=True)

        scrollbar = ttk.Scrollbar(table_frame, orient=VERTICAL, command=self.tree.yview)
        self.tree.configure(yscrollcommand=scrollbar.set)
        scrollbar.pack(side="right", fill="y")

        # Editable cells
        self.tree.bind('<Double-1>', self.on_double_click)

        # --- Note below the table ---
        note_label = Label(main_frame,
                              text="Note: The 'Disbursal Date' will be set to the 'Month Year' entered above when creating the invoice.",
                              font=('Arial', 9, 'italic'),
                              fg='gray')
        note_label.pack(pady=(5, 0))

        # --- Buttons ---
        button_frame = Frame(main_frame)
        button_frame.pack(fill="x", pady=5)

        self.create_button = Button(button_frame, text="Create Invoice", command=self.create_invoice)
        self.create_button.pack(side="left", padx=5, expand=True)

        # Load Excel data
        self.load_data_from_excel()

    @staticmethod
    def IntComma(num):
        parts = num.split('.')
        integer_part = parts[0]
        fractional_part = parts[1] if len(parts) > 1 else None

        # Format the integer part with commas
        n = len(integer_part)
        if n <= 3:
            formatted_integer_part = integer_part
        else:
            result = []
            for i in range(n - 1, -1, -3):
                start = max(0, i - 2)
                result.append(integer_part[start:i + 1])
            formatted_integer_part = ",".join(reversed(result))

        # Reassemble the number
        if fractional_part is not None:
            return f"{formatted_integer_part}.{fractional_part}"
        else:
            return formatted_integer_part


    @staticmethod
    def clean_and_convert_Integer(raw_string):
        # Regex: Find and remove everything that is NOT a digit (\d), a dot (\.), or a minus sign (\-)
        # This single line handles spaces, commas, and \xa0 characters effectively.
        cleaned_string = sub(r'[^\d\.\-]', '', raw_string)
        
        # The sub function alone is often enough, but a final .strip() is safe practice
        return str(cleaned_string.strip())
    
    @staticmethod
    def clean_and_convert_String(raw_string):
        # Regex: Find and remove everything that is NOT a digit (\d), a dot (\.), or a minus sign (\-)
        # This single line handles spaces, commas, and \xa0 characters effectively.
        cleaned_string = raw_string.replace('\xa0', '').strip()
        
        # The sub function alone is often enough, but a final .strip() is safe practice
        return str(cleaned_string.strip())


    def load_data_from_excel(self):
        """Loads customer data from the Excel file into the Treeview."""
        try:
            for item in self.tree.get_children():
                self.tree.delete(item)

            df = read_excel(self.excel_file_path, header=None)

            # TODO: FIX PAYMENT SLAB ISSUE AND REFER TO TABULAR FORMAT OF DIB
            # Columns: 1, 2, 3, 6, 7, 8, 9 in Excel. Payout and VAT are calculated.
            self.customer_data = df.iloc[:, [5, 1, 2, 3, 6, 7]].copy()

            for _, row_data in self.customer_data.iterrows():
                try:
                    # Clean and process data from Excel
                    disbursal_date_str = "(Month Year as stated above)"
                    app_ref = self.clean_and_convert_String(row_data.iloc[1]) if notna(row_data.iloc[1]) else ""
                    customer_name = self.clean_and_convert_String(row_data.iloc[2]).title() if notna(row_data.iloc[2]) else ""
                    
                    # Clean all numeric strings using sub
                    loan_amount = float(self.clean_and_convert_Integer(row_data.iloc[3]))
                
                    payment_slab = float(self.clean_and_convert_String(str(row_data.iloc[4]))) * 100
                    
                    # Calculated values
                    Incentive = loan_amount * (payment_slab / 100)
                    payout = Incentive / 1.05
                    fivePercentVat = Incentive / 21

                    self.tree.insert("", END, values=(
                        disbursal_date_str, app_ref, customer_name,
                        f"{loan_amount:.2f}", f"{payment_slab:.2f}%",
                        f"{payout:.2f}", f"{fivePercentVat:.2f}", f"{Incentive:.2f}"
                    ))

                except Exception as e:
                    # TODO: DELETE STMT DURING DEPLOYMENT
                    print(f"Error processing row: {e}")
                    pass
            if not self.tree.get_children():
                raise ValueError("No valid rows loaded from Excel")

        except FileNotFoundError:
            messagebox.showerror("Error", f"Excel file not found at: {self.excel_file_path}")

        except IndexError:
            messagebox.showerror("Error", "Excel file does not contain the required columns.\nVerify correct File / Format(.xlsx) and it's columns.")
            self.root.destroy()
            self.main_app_root.destroy()

        except ValueError as e:
            messagebox.showerror("Error", f"Verify Correct File / Format(.xlsx) and it's columns.\nError: {e}")
            self.root.destroy()
            self.main_app_root.destroy()

        except TclError as e:
            messagebox.showerror("Error", f"Failed to load Excel data.\nVerify Correct File / Format(.xlsx) and it's columns.\nError: {e}")
            self.root.destroy()
            self.main_app_root.destroy()

        except Exception as e:
            messagebox.showerror("Error", f"Unknown Error: {e}.\nContact the developer (Aryan).")

    def on_double_click(self, event):
        # Enable inline editing for cells, except Incentive, VAT, and Payout
        region = self.tree.identify("region", event.x, event.y)
        if region != "cell":
            return

        column = self.tree.identify_column(event.x)
        column_index = int(column[1:]) - 1  # 0-based

        if column_index in (5, 6, 7):  # Payout, VAT, or Incentive column
            messagebox.showinfo("Info", "This field is automatically calculated and not editable.")
            return

        item_id = self.tree.identify_row(event.y)
        if not item_id:
            return

        cell_value = self.tree.item(item_id, "values")[column_index]
        x, y, width, height = self.tree.bbox(item_id, column)

        entry_edit = ttk.Entry(self.tree, justify="center")
        entry_edit.place(x=x, y=y, width=width, height=height)
        entry_edit.insert(0, cell_value)
        entry_edit.focus()

        def save_edit(event):
            new_value = entry_edit.get()
            current_values = list(self.tree.item(item_id, "values"))
            current_values[column_index] = new_value

            # Recalculate Payout, VAT and Incentive if Loan Amount or Payment Slab changed
            if column_index in (3, 4):
                try:
                    loan_amount_str = sub(r'[^\d.]', '', str(current_values[3]))
                    loan_amount = float(loan_amount_str)

                    payment_slab_str = sub(r'[^\d.]', '', str(current_values[4]))
                    payment_slab = float(payment_slab_str)
                
                    current_values[4] = str(payment_slab)
                    
                    payout_calculated = loan_amount * payment_slab / 100
                    vat_calculated = payout_calculated * 0.05
                    incentive_calculated = payout_calculated + vat_calculated

                    current_values[4] += "%" if "%" not in current_values[4] else ""
                    
                    current_values[5] = self.IntComma(f"{payout_calculated:.2f}")
                    current_values[6] = self.IntComma(f"{vat_calculated:.2f}")
                    current_values[7] = self.IntComma(f"{incentive_calculated:.2f}")
                    
                except ValueError:
                    messagebox.showerror("Invalid Input", "Loan Amount and Payment Slab must be numeric.")
                    return

            self.tree.item(item_id, values=current_values)
            entry_edit.destroy()

        entry_edit.bind("<Return>", save_edit)
        entry_edit.bind("<FocusOut>", save_edit)

    def run(self):
        self.root.mainloop()

    def on_close(self):
        self.root.destroy()
        self.main_app_root.destroy()

    @staticmethod
    def replace_text(paragraph, old_text, new_text, isTable=False):
        full = "".join(run.text for run in paragraph.runs) or paragraph.text
        if old_text in full:
            replaced = full.replace(old_text, new_text)
            if paragraph.runs:
                paragraph.runs[0].text = replaced
                for r in paragraph.runs[1:]:
                    r.text = ""
            else:
                paragraph.text = replaced

    def format_table_cells(self, table):
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
    
    @staticmethod
    def number_to_words(num):
        try:
            # Handle integers
            if isinstance(num, int):
                return num2words(num)

            # Handle floats
            elif isinstance(num, float):
                integer_part = int(num)
                decimal_part = str(num).split(".")[1]

                # Convert integer part
                wordsFirst = ''.join(num2words(integer_part).split(" and"))
                wordsSecond = num2words(decimal_part)

                return f"{wordsFirst}".title() + " and " + f"{wordsSecond} Fills Only".title()

            else:
                return "Invalid Number"

        except Exception as e:
            return f"Error: {e}"
        
    @staticmethod
    def convertToFull(month_year):
        try:
            
            if " " in month_year:
                month_year = month_year.split(" ")
                month = month_year[0]
                year = month_year[1]

            else :
                month = ""
                other = ""
                year = ""

                for i in month_year:
                    if i.isalpha():
                        month += i
                    elif i.isnumeric():
                        year += i
                    else:
                        other += i

            months = {
                "jan": "January",
                "feb": "February",
                "mar": "March",
                "apr": "April",
                "may": "May",
                "jun": "June",
                "jul": "July",
                "aug": "August",
                "sep": "September",
                "oct": "October",
                "nov": "November",
                "dec": "December"
            }

            return months.get(month, month) + ((" " + year) if year else year)
        except IndexError:
            messagebox.showerror("Missing Information", "Please enter month and year")

    def create_invoice(self):
        """Generate invoice Word + PDF with table filled from Treeview."""
        if not self.invoice_number_entry.get():
            messagebox.showerror("Missing Information", "Please enter the Invoice Number.")
            return

        if not self.month_year_entry.get():
            messagebox.showerror("Missing Information", "Please enter the Month Year.")
            return

        items = self.tree.get_children()
        if not items:
            messagebox.showerror("Error", "No data loaded.")
            return
        
        self.root.withdraw()  # Hide the current window

        totalLoanAmount = 0.0
        totalPayout = 0.0
        totalVAT = 0.0
        totalIncentive = 0.0

        doc = Document(self.template)

        replacements = {
            "[date today]": datetime.today().strftime('%d/%m/%Y'),
            "[invoice number]": self.invoice_number_entry.get(),
            "[FullMonth year]" : self.convertToFull(self.month_year_entry.get().lower()).title(),
            "[month year]": self.month_year_entry.get().title()
        }

        customer_table = None
        for table in doc.tables:
            headers = (cell.text.strip().lower() for cell in table.rows[0].cells)
            if "customer name" in headers and "loan amount" in headers:
                customer_table = table
                break

        if not customer_table:
            messagebox.showerror("Error", "Customer table not found in template.")
            return

        for item_id in items:
            row_data = self.tree.item(item_id, "values")
            try:
                row = customer_table.add_row().cells

                # Disbursal Date
                row[0].text = self.month_year_entry.get().title() or str(row_data[0])

                # App reference
                row[1].text = str(row_data[1])
                row[1].paragraphs[0].alignment = enum.text.WD_ALIGN_PARAGRAPH.CENTER

                # Customer Name
                row[2].text = str(row_data[2])

                # Loan Amount
                loan_amount_str = sub(r'[^\d.]', '', str(row_data[3]))
                loan_amount = float(loan_amount_str)
                row[3].text = self.IntComma(f"{loan_amount:.2f}")
                row[3].paragraphs[0].alignment = enum.text.WD_ALIGN_PARAGRAPH.CENTER
                totalLoanAmount += loan_amount

                # Payment Slab
                row[4].text = str(row_data[4])
                row[4].paragraphs[0].alignment = enum.text.WD_ALIGN_PARAGRAPH.CENTER

                # Payout
                payout_str = sub(r'[^\d.]', '', str(row_data[5]))
                payout = float(payout_str)
                row[5].text = self.IntComma(f"{payout:.2f}")
                row[5].paragraphs[0].alignment = enum.text.WD_ALIGN_PARAGRAPH.CENTER
                totalPayout += payout

                # VAT
                vat_str = sub(r'[^\d.]', '', str(row_data[6]))
                vat = float(vat_str)
                row[6].text = self.IntComma(f"{vat:.2f}")
                row[6].paragraphs[0].alignment = enum.text.WD_ALIGN_PARAGRAPH.CENTER
                totalVAT += vat

                # Incentive
                incentive_str = sub(r'[^\d.]', '', str(row_data[7]))
                incentive = float(incentive_str)
                row[7].text = self.IntComma(f"{incentive:.2f}")
                row[7].paragraphs[0].alignment = enum.text.WD_ALIGN_PARAGRAPH.CENTER
                totalIncentive += incentive

            except Exception as e:
                messagebox.showerror("Data Error", f"Row error: {e}")
                return

        replacements.update({
            "[total loan]": self.IntComma(f"{totalLoanAmount:.2f}"),
            "[total payout]": self.IntComma(f"{totalPayout:.2f}"),
            "[total vat]": self.IntComma(f"{totalVAT:.2f}"),
            "[VAT&incent]": self.IntComma(f"{totalIncentive:.2f}"),
            "[AmtinWords]" : f"{self.number_to_words(totalIncentive)}"
        })

        self.format_table_cells(customer_table)

        for paragraph in doc.paragraphs:
            for old_text, new_text in replacements.items():
                self.replace_text(paragraph, old_text, new_text)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for old_text, new_text in replacements.items():
                            self.replace_text(paragraph, old_text, new_text)

        try:
            save_path = filedialog.asksaveasfilename(defaultextension=".pdf", filetypes=[("PDF Documents", "*.pdf")])
            if save_path:
                doc.save('filled.docx')
                convert('filled.docx', save_path)
                messagebox.showinfo("Success", "Invoice created and saved successfully!")
            try:
                self.root.destroy()
                self.main_app_root.destroy()
            except TclError:
                # THE WINDOW HAS ALREADY BEEN DESTROYED
                pass

            # IGNORE ANY OTHER EXCEPTION
            except Exception as e:
                pass
        except PermissionError:
            messagebox.showerror("Error", "Permission denied. Close the file if it's open and try again.")